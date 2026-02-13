from docx import Document

def export_docx(sections, path="report.docx"):
    doc = Document()
    doc.add_heading("Account Strategy Research Report", 0)

    for name, data in sections.items():
        doc.add_heading(name, level=1)
        doc.add_paragraph(data["content"])

        doc.add_heading("References", level=2)
        for c in data["citations"]:
            p = doc.add_paragraph()
            run = p.add_run(f"[{c['id']}] {c['title']} ({c['url']})")
            run.font.underline = True

        doc.add_paragraph(
            f"Confidence Score: {round(data['confidence']*100)}%"
        )

    doc.save(path)
    return path
