import os
import json
import operator
import asyncio
from typing import Annotated, List, TypedDict
from fastapi import FastAPI, Form, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from langchain_anthropic import ChatAnthropic
from langchain_community.tools.tavily_search import TavilySearchResults
from langgraph.graph import StateGraph, END
from langgraph.constants import Send

# Formatting Imports
import markdown2
from weasyprint import HTML
from docx import Document # Requires: pip install python-docx

# --- INITIALIZATION ---
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

llm = ChatAnthropic(model="claude-3-5-sonnet-latest")
search_tool = TavilySearchResults(max_results=5)

# --- LANGGRAPH STATE DEFINITIONS ---
class SectionState(TypedDict):
    section_name: str
    target_company: str

class OverallState(TypedDict):
    target_company: str
    research_data: Annotated[List[dict], operator.add]
    final_report: str

# --- GRAPH NODES ---
def planner(state: OverallState):
    sections = [
        "Account Business Overview",
        "Key Business Initiatives",
        "Tech Landscape",
        "Strategic Next Steps"
    ]
    return [Send("researcher", {"section_name": s, "target_company": state['target_company']}) for s in sections]

def researcher(state: SectionState):
    query = f"{state['target_company']} {state['section_name']} news 2024 2025"
    search_results = search_tool.invoke(query)
    prompt = f"Summarize research for {state['section_name']} of {state['target_company']}. Results: {search_results}"
    response = llm.invoke(prompt)
    return {"research_data": [{"section": state['section_name'], "content": response.content}]}

def writer(state: OverallState):
    report = f"# Strategic Report: {state['target_company']}\n\n"
    for sec in sorted(state['research_data'], key=lambda x: x['section']):
        report += f"## {sec['section']}\n{sec['content']}\n\n"
    return {"final_report": report}

# --- BUILD THE GRAPH ---
builder = StateGraph(OverallState)
builder.add_node("planner", planner)
builder.add_node("researcher", researcher)
builder.add_node("writer", writer)
builder.set_entry_point("planner")
builder.add_conditional_edges("planner", lambda x: x)
builder.add_edge("researcher", "writer")
builder.add_edge("writer", END)
graph = builder.compile()

# --- API ENDPOINTS ---
reports_db = {}

@app.post("/research")
async def run_research(target_company: str = Form(...), annual_report: UploadFile = None):
    async def stream_generator():
        yield json.dumps({"type": "status", "message": "Analyzing Data..."}) + "\n"
        state = {"target_company": target_company, "research_data": []}
        async for event in graph.astream(state):
            for node, output in event.items():
                if node == "researcher":
                    section_info = output["research_data"][0]
                    yield json.dumps({
                        "type": "section",
                        "data": {"section_title": section_info["section"], "section_content": section_info["content"]}
                    }) + "\n"
                elif node == "writer":
                    report_id = target_company.replace(" ", "_").lower()
                    reports_db[report_id] = output["final_report"]
        yield json.dumps({"type": "complete"}) + "\n"
    return StreamingResponse(stream_generator(), media_type="text/event-stream")

# --- EXPORT ENDPOINTS ---

@app.get("/download/pdf/{report_id}")
async def download_pdf(report_id: str):
    content = reports_db.get(report_id)
    if not content: raise HTTPException(status_code=404)
    pdf_path = f"/tmp/{report_id}.pdf"
    HTML(string=markdown2.markdown(content)).write_pdf(pdf_path)
    return FileResponse(pdf_path, filename=f"Research_{report_id}.pdf")

@app.get("/download/docx/{report_id}")
async def download_docx(report_id: str):
    content = reports_db.get(report_id)
    if not content: raise HTTPException(status_code=404)
    
    doc_path = f"/tmp/{report_id}.docx"
    doc = Document()
    
    # Simple Markdown-to-Docx conversion
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        else:
            doc.add_paragraph(line)
            
    doc.save(doc_path)
    return FileResponse(doc_path, filename=f"Research_{report_id}.docx")
